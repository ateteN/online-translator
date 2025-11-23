from fastapi import FastAPI, UploadFile, Form
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from docx import Document
import fitz
import openai
import uvicorn
import os

openai.api_key = os.getenv("OPENAI_API_KEY")

app = FastAPI()

# Allow frontend domain
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# --- DOCX READER ---
def read_docx(path):
    doc = Document(path)
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

# --- PDF READER ---
def read_pdf(path):
    pdf = fitz.open(path)
    paragraphs = []
    for page in pdf:
        text = page.get_text().strip()
        for p in text.split("\n\n"):
            if p.strip():
                paragraphs.append(p.strip())
    return paragraphs

# --- TRANSLATION ---
def translate(text, target_lang):
    prompt = f"""
Translate this paragraph into {target_lang}:

Original:
{text}

Translation:
"""
    response = openai.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message['content'].strip()

# --- DOCX OUTPUT ---
def write_doc(paragraphs, translations, output_path):
    doc = Document()
    for o, t in zip(paragraphs, translations):
        doc.add_paragraph(o)
        doc.add_paragraph(t)
        doc.add_paragraph("")
    doc.save(output_path)

@app.post("/translate")
async def translate_doc(document: UploadFile, language: str = Form(...)):
    filename = "input." + document.filename.split(".")[-1]
    
    with open(filename, "wb") as f:
        f.write(await document.read())

    if filename.endswith(".docx"):
        paragraphs = read_docx(filename)
    else:
        paragraphs = read_pdf(filename)

    translations = [translate(p, language) for p in paragraphs]

    output_path = "translated_output.docx"
    write_doc(paragraphs, translations, output_path)

    return FileResponse(output_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

